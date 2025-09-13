from django.shortcuts import render
from django.conf import settings

def debug_template_settings(request):
    from pprint import pformat
    return HttpResponse(
        "BASE_DIR:\n{}\n\nTEMPLATE DIRS:\n{}\n\nINSTALLED_APPS:\n{}".format(
            settings.BASE_DIR,
            pformat(settings.TEMPLATES[0]['DIRS']),
            pformat(settings.INSTALLED_APPS)
        ),
        content_type="text/plain"
    )

from income.models import Income
from expenditure.models import Expenditure
from savings.models import Saving
from django.db import models

from django.contrib.auth.decorators import login_required

@login_required
def home(request):
    total_income = Income.objects.filter(user=request.user).aggregate(total=models.Sum("amount"))["total"] or 0
    total_expenditure = Expenditure.objects.filter(user=request.user).aggregate(total=models.Sum("amount"))["total"] or 0
    total_savings = Saving.objects.filter(user=request.user).aggregate(total=models.Sum("amount"))["total"] or 0
    context = {
        "total_income": total_income,
        "total_expenditure": total_expenditure,
        "total_savings": total_savings
    }
    return render(request, "home.html", context)

from django.http import HttpResponse

import pandas as pd
from io import BytesIO

@login_required
def export_report_excel(request):
    incomes = Income.objects.values("date", "amount", "source", "description")
    expenditures = Expenditure.objects.values("date", "amount", "category", "description")
    savings = Saving.objects.values("date", "amount", "target", "description")

    df_incomes = pd.DataFrame(list(incomes))
    df_expenditures = pd.DataFrame(list(expenditures))
    df_savings = pd.DataFrame(list(savings))

    with BytesIO() as b:
        with pd.ExcelWriter(b, engine="openpyxl") as writer:
            df_incomes.to_excel(writer, sheet_name="Income", index=False)
            df_expenditures.to_excel(writer, sheet_name="Expenditure", index=False)
            df_savings.to_excel(writer, sheet_name="Savings", index=False)
        b.seek(0)
        response = HttpResponse(
            b.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = "attachment; filename=financial_report.xlsx"
        return response

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

@login_required
def export_report_pdf(request):
    incomes = Income.objects.values_list("date", "amount", "source", "description")
    expenditures = Expenditure.objects.values_list("date", "amount", "category", "description")
    savings = Saving.objects.values_list("date", "amount", "target", "description")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)

    elements = []
    styles = getSampleStyleSheet()
    title = Paragraph("Financial Report", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))

    # Income Section
    elements.append(Paragraph("Income Records", styles['Heading2']))
    inc_data = [["Date", "Amount", "Source", "Description"]] + list(incomes)
    inc_table = Table(inc_data, hAlign='LEFT')
    inc_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    elements.append(inc_table)
    elements.append(Spacer(1, 18))

    # Expenditure Section
    elements.append(Paragraph("Expenditure Records", styles['Heading2']))
    exp_data = [["Date", "Amount", "Category", "Description"]] + list(expenditures)
    exp_table = Table(exp_data, hAlign='LEFT')
    exp_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.salmon),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    elements.append(exp_table)
    elements.append(Spacer(1, 18))

    # Savings Section
    elements.append(Paragraph("Savings Records", styles['Heading2']))
    sav_data = [["Date", "Amount", "Target", "Description"]] + list(savings)
    sav_table = Table(sav_data, hAlign='LEFT')
    sav_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgreen),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    elements.append(sav_table)

    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=financial_report.pdf"
    return response

from docx import Document
from docx.shared import Inches
from income.forms import RegistrationForm
from django.contrib.auth import login as auth_login

@login_required
def export_report_word(request):
    incomes = Income.objects.values_list("date", "amount", "source", "description")
    expenditures = Expenditure.objects.values_list("date", "amount", "category", "description")
    savings = Saving.objects.values_list("date", "amount", "target", "description")

    doc = Document()
    doc.add_heading('Financial Report', 0)

    # Income Section
    doc.add_heading('Income Records', level=1)
    inc_table = doc.add_table(rows=1, cols=4)
    hdr_cells = inc_table.rows[0].cells
    hdrs = ["Date", "Amount", "Source", "Description"]
    for i, h in enumerate(hdrs):
        hdr_cells[i].text = h
    for row in incomes:
        row_cells = inc_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_paragraph('')

    # Expenditure Section
    doc.add_heading('Expenditure Records', level=1)
    exp_table = doc.add_table(rows=1, cols=4)
    hdr_cells = exp_table.rows[0].cells
    hdrs = ["Date", "Amount", "Category", "Description"]
    for i, h in enumerate(hdrs):
        hdr_cells[i].text = h
    for row in expenditures:
        row_cells = exp_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_paragraph('')

    # Savings Section
    doc.add_heading('Savings Records', level=1)
    sav_table = doc.add_table(rows=1, cols=4)
    hdr_cells = sav_table.rows[0].cells
    hdrs = ["Date", "Amount", "Target", "Description"]
    for i, h in enumerate(hdrs):
        hdr_cells[i].text = h
    for row in savings:
        row_cells = sav_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # Output  
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=financial_report.docx"
    return response

def register(request):
    if request.method == "POST":
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            return render(request, "register_done.html")
        # falls through to redisplay with errors
    else:
        form = RegistrationForm()
    return render(request, "register.html", {"form": form})
